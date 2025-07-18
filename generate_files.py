import os
import random
from faker import Faker
from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

fake = Faker()


def generate_text(min_sentences=3, max_sentences=15):
    """Генерация разнообразного текстового контента"""
    num_sentences = random.randint(min_sentences, max_sentences)
    content_types = [
        lambda: fake.paragraph(nb_sentences=random.randint(2, 5)),
        lambda: fake.text(max_nb_chars=random.randint(200, 800)),
        lambda: '\n'.join(f'- {fake.sentence()}' for _ in range(random.randint(3, 7))),
        lambda: '\n'.join(f'{i + 1}. {fake.sentence()}' for i in range(random.randint(3, 7))),
        lambda: f"{fake.catch_phrase().upper()}\n\n{fake.paragraph(nb_sentences=4)}",
        lambda: f"{fake.company()} Report\n\n{fake.paragraph(nb_sentences=3)}",
        lambda: f"Date: {fake.date_this_decade()}\n\nSubject: {fake.sentence()}\n\n{fake.paragraph(nb_sentences=4)}",
        lambda: f"{fake.word().capitalize()} Analysis\n\n{fake.paragraph(nb_sentences=5)}",
        lambda: f"CONTACT:\nName: {fake.name()}\nEmail: {fake.email()}\nPhone: {fake.phone_number()}\nAddress: {fake.address().replace('\n', ', ')}"
    ]
    return '\n\n'.join(random.choice(content_types)() for _ in range(num_sentences))


def generate_table(rows, cols):
    """Генерация таблиц с различными типами данных"""
    table_types = ["financial", "personnel", "technical", "mixed"]
    table_type = random.choice(table_types)

    if table_type == "financial":
        return [
            ['Period', 'Revenue', 'Expenses', 'Profit', 'Growth'],
            *[
                [
                    f"Q{i + 1} {random.randint(2020, 2023)}",
                    f"${random.randint(10000, 1000000):,.2f}",
                    f"${random.randint(5000, 800000):,.2f}",
                    f"${random.randint(5000, 500000):,.2f}",
                    f"{random.uniform(-5.0, 25.0):.2f}%"
                ] for i in range(rows - 1)
            ]
        ]

    elif table_type == "personnel":
        return [
            ['ID', 'Name', 'Position', 'Department', 'Salary', 'Hire Date'],
            *[
                [
                    f"EMP{random.randint(1000, 9999)}",
                    fake.name(),
                    fake.job(),
                    fake.bs(),
                    f"${random.randint(30000, 150000):,}",
                    fake.date_this_decade().strftime('%Y-%m-%d')
                ] for _ in range(rows - 1)
            ]
        ]

    elif table_type == "technical":
        components = ['CPU', 'GPU', 'RAM', 'SSD', 'Motherboard', 'PSU']
        return [
            ['Component', 'Model', 'Specs', 'Qty', 'Price'],
            *[
                [
                    random.choice(components),
                    fake.bothify(text='??-####'),
                    f"{random.choice(['4', '8', '16'])}GB {random.choice(['DDR4', 'DDR5'])}",
                    random.randint(1, 10),
                    f"${random.randint(50, 500):.2f}"
                ] for _ in range(rows - 1)
            ]
        ]

    else:  # mixed
        return [
            [
                random.choice([
                    fake.word().capitalize(),
                    str(random.randint(1, 10000)),
                    fake.date_this_decade().strftime('%Y-%m-%d'),
                    random.choice(['Yes', 'No']),
                    f"=A{i + 1}*{random.uniform(0.5, 2.0):.2f}"
                ]) for j in range(cols)
            ] for i in range(rows)
        ]


def generate_docx(pages, file_path):
    """Генерация DOCX файла"""
    doc = Document()
    for page_num in range(1, pages + 1):
        # Заголовок
        doc.add_heading(f"{fake.catch_phrase()} - Page {page_num}", level=0)

        # Текстовый контент
        for _ in range(random.randint(2, 4)):
            doc.add_paragraph(generate_text(2, 5))

        # Подзаголовок
        doc.add_heading(fake.catch_phrase(), level=2)

        # Список
        list_type = random.choice(['bullet', 'number'])
        list_items = [fake.sentence() for _ in range(random.randint(3, 6))]
        for item in list_items:
            if list_type == 'bullet':
                doc.add_paragraph(item, style='ListBullet')
            else:
                doc.add_paragraph(item, style='ListNumber')

        # Таблица
        table = doc.add_table(rows=random.randint(4, 8), cols=random.randint(3, 5))
        table.style = 'Table Grid'
        table_data = generate_table(len(table.rows), len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.text = str(table_data[i][j])

        # Разрыв страницы
        if page_num < pages:
            doc.add_page_break()

    doc.save(file_path)


def generate_pdf(pages, file_path):
    """Генерация PDF файла"""
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(file_path, pagesize=A4)
    elements = []

    for page_num in range(1, pages + 1):
        # Заголовок
        elements.append(Paragraph(f"{fake.catch_phrase()} - Page {page_num}", styles['Heading1']))
        elements.append(Spacer(1, 12))

        # Текстовый контент
        for _ in range(random.randint(2, 4)):
            elements.append(Paragraph(generate_text(2, 4), styles['BodyText']))
            elements.append(Spacer(1, 6))

        # Подзаголовок
        elements.append(Paragraph(fake.catch_phrase(), styles['Heading2']))
        elements.append(Spacer(1, 6))

        # Список
        list_type = random.choice(['bullet', 'number'])
        list_items = [fake.sentence() for _ in range(random.randint(3, 5))]
        for item in list_items:
            if list_type == 'bullet':
                elements.append(Paragraph(f"• {item}", styles['BodyText']))
            else:
                elements.append(Paragraph(f"{list_items.index(item) + 1}. {item}", styles['BodyText']))

        elements.append(Spacer(1, 12))

        # Таблица
        table_data = generate_table(random.randint(4, 8), random.randint(3, 5))
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(table)

        # Разрыв страницы
        if page_num < pages:
            elements.append(Spacer(1, 24))

    doc.build(elements)


def generate_xlsx(pages, file_path):
    """Генерация XLSX файла"""
    wb = Workbook()
    wb.remove(wb.active)  # Удаляем дефолтный лист

    for page_num in range(1, pages + 1):
        ws = wb.create_sheet(title=f"Page_{page_num}")

        # Заголовок
        ws['A1'] = f"{fake.company()} - Report"

        # Таблицы
        for table_num in range(1, random.randint(2, 3)):
            start_row = (table_num - 1) * 15 + 3
            table_data = generate_table(random.randint(5, 10), random.randint(4, 6))

            # Запись данных таблицы
            for i, row in enumerate(table_data):
                for j, value in enumerate(row):
                    ws.cell(row=start_row + i, column=j + 1, value=value)

        # Текстовый блок
        ws.cell(row=start_row + len(table_data) + 2, column=1, value="Summary")
        ws.cell(row=start_row + len(table_data) + 3, column=1, value=generate_text(2, 3))

    wb.save(file_path)


def generate_all_files():
    """Генерация всех тестовых файлов (в 2 раза больше, без указания размера в имени)"""
    sizes = [1, 5, 20]
    os.makedirs("test_files", exist_ok=True)

    # Счетчики для каждого типа файлов
    counters = {
        'docx': 1,
        'pdf': 1,
        'xlsx': 1
    }

    for size in sizes:
        for _ in range(2):  # 2 файла каждого типа и размера
            # DOCX
            generate_docx(size, f"test_files/docx_{counters['docx']}.docx")
            counters['docx'] += 1

            # PDF
            generate_pdf(size, f"test_files/pdf_{counters['pdf']}.pdf")
            counters['pdf'] += 1

            # XLSX
            generate_xlsx(size, f"test_files/xlsx_{counters['xlsx']}.xlsx")
            counters['xlsx'] += 1


if __name__ == "__main__":
    generate_all_files()